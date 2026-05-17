import streamlit as st
from PIL import Image, ImageOps
import io
import base64
import math
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas as rl_canvas
from docx import Document
from docx.shared import Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn as docx_qn
from docx.oxml import OxmlElement as DocxEl

st.set_page_config(page_title="Image to A4 Organizer", layout="wide")



# ── helpers ────────────────────────────────────────────────────────────────────

def fix_orientation(img):
    try:
        return ImageOps.exif_transpose(img)
    except Exception:
        return img


def strip_metadata(img):
    clean = img.copy()
    clean.info = {}
    return clean


def compress_image(img, level):
    quality = {"Low": 82, "Medium": 55, "High": 22}[level]
    buf = io.BytesIO()
    img.convert("RGB").save(buf, "JPEG", quality=quality, optimize=True)
    buf.seek(0)
    return Image.open(buf).copy()


def to_jpeg_buf(img, quality=92):
    buf = io.BytesIO()
    (img if img.mode == "RGB" else img.convert("RGB")).save(buf, "JPEG", quality=quality)
    buf.seek(0)
    return buf


def grid_shape(ipp):
    if ipp == 1:
        return 1, 1
    if ipp == 2:
        return 2, 1
    cols = math.ceil(math.sqrt(ipp))
    rows = math.ceil(ipp / cols)
    return cols, rows


# ── generators ─────────────────────────────────────────────────────────────────

def generate_pdf(images, ipp, comp, do_strip=True):
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=A4)
    pw, ph = A4
    margin = 28.35   # 10 mm in points
    pad = 6          # inner padding per cell

    for start in range(0, len(images), ipp):
        batch = images[start : start + ipp]
        cols, rows = grid_shape(ipp)
        cw = (pw - 2 * margin) / cols
        ch = (ph - 2 * margin) / rows

        for idx, img in enumerate(batch):
            img = fix_orientation(img)
            if do_strip:
                img = strip_metadata(img)
            if comp:
                img = compress_image(img, comp)

            col, row = idx % cols, idx // cols
            aw, ah = cw - 2 * pad, ch - 2 * pad
            iw, ih = img.size
            scale = min(aw / iw, ah / ih)
            dw, dh = iw * scale, ih * scale
            x = margin + col * cw + (cw - dw) / 2
            y = ph - margin - (row + 1) * ch + (ch - dh) / 2

            c.drawImage(ImageReader(to_jpeg_buf(img)), x, y, dw, dh,
                        preserveAspectRatio=True, mask="auto")

        c.showPage()

    c.save()
    buf.seek(0)
    return buf.read()


def generate_html(images, ipp, comp, do_strip=True):
    cols, rows = grid_shape(ipp)
    pages = []

    for start in range(0, len(images), ipp):
        batch = images[start : start + ipp]
        cells = []
        for img in batch:
            img = fix_orientation(img)
            if do_strip:
                img = strip_metadata(img)
            if comp:
                img = compress_image(img, comp)
            b64 = base64.b64encode(to_jpeg_buf(img).read()).decode()
            cells.append(
                f'<div class="cell">'
                f'<img src="data:image/jpeg;base64,{b64}" alt="">'
                f"</div>"
            )
        pages.append(
            f'<div class="page"><div class="grid">{"".join(cells)}</div></div>'
        )

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Images</title>
<style>
  *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ background: #555; }}
  .page {{
    width: 210mm; height: 297mm;
    background: #fff;
    margin: 8mm auto;
    padding: 8mm;
    page-break-after: always;
    break-after: page;
    display: flex;
  }}
  .grid {{
    width: 100%; height: 100%;
    display: grid;
    grid-template-columns: repeat({cols}, 1fr);
    grid-template-rows: repeat({rows}, 1fr);
    gap: 4mm;
  }}
  .cell {{ display: flex; align-items: center; justify-content: center; overflow: hidden; }}
  .cell img {{ max-width: 100%; max-height: 100%; object-fit: contain; display: block; }}
  @media print {{
    body {{ background: #fff; }}
    .page {{ margin: 0; box-shadow: none; }}
    @page {{ size: A4; margin: 0; }}
  }}
</style>
</head>
<body>
{"".join(pages)}
</body>
</html>"""


def generate_docx(images, ipp, comp, do_strip=True):
    def _no_borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = DocxEl("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = DocxEl(f"w:{side}")
            b.set(docx_qn("w:val"), "none")
            b.set(docx_qn("w:sz"), "0")
            b.set(docx_qn("w:space"), "0")
            b.set(docx_qn("w:color"), "auto")
            borders.append(b)
        tblPr.append(borders)

    def _cell_width(cell, mm):
        tcPr = cell._tc.get_or_add_tcPr()
        tcW = DocxEl("w:tcW")
        tcW.set(docx_qn("w:w"), str(int(mm * 56.6929)))
        tcW.set(docx_qn("w:type"), "dxa")
        tcPr.append(tcW)

    def _row_height(row, mm):
        trPr = row._tr.get_or_add_trPr()
        trH = DocxEl("w:trHeight")
        trH.set(docx_qn("w:val"), str(int(mm * 56.6929)))
        trH.set(docx_qn("w:hRule"), "exact")
        trPr.append(trH)

    def _vcenter(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        v = DocxEl("w:vAlign")
        v.set(docx_qn("w:val"), "center")
        tcPr.append(v)

    doc = Document()
    # Remove the default empty paragraph Word always adds
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    m = 10  # margin mm
    sec.top_margin = Mm(m)
    sec.bottom_margin = Mm(m)
    sec.left_margin = Mm(m)
    sec.right_margin = Mm(m)

    cols_n, rows_n = grid_shape(ipp)
    gap = 4  # mm gap between cells
    avail_w = 210 - 2 * m   # 190 mm
    avail_h = 297 - 2 * m   # 277 mm
    cell_w_mm = (avail_w - gap * (cols_n - 1)) / cols_n
    cell_h_mm = (avail_h - gap * (rows_n - 1)) / rows_n
    img_w_mm = cell_w_mm - 2
    img_h_mm = cell_h_mm - 2

    for page_idx, start in enumerate(range(0, len(images), ipp)):
        if page_idx > 0:
            pb_para = doc.add_paragraph()
            pb_para.paragraph_format.space_before = Mm(0)
            pb_para.paragraph_format.space_after = Mm(0)
            br = DocxEl("w:br")
            br.set(docx_qn("w:type"), "page")
            pb_para.add_run()._r.append(br)

        batch = images[start : start + ipp]
        table = doc.add_table(rows=rows_n, cols=cols_n)
        _no_borders(table)

        for row in table.rows:
            _row_height(row, cell_h_mm)
            for cell in row.cells:
                _cell_width(cell, cell_w_mm)
                _vcenter(cell)

        for idx, img in enumerate(batch):
            img = fix_orientation(img)
            if do_strip:
                img = strip_metadata(img)
            if comp:
                img = compress_image(img, comp)
            col, row = idx % cols_n, idx // cols_n
            cell = table.cell(row, col)

            iw, ih = img.size
            max_w_in = img_w_mm / 25.4
            max_h_in = img_h_mm / 25.4
            fit_w = max_w_in
            fit_h = fit_w * ih / iw
            if fit_h > max_h_in:
                fit_h = max_h_in
                fit_w = fit_h * iw / ih

            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Mm(0)
            para.paragraph_format.space_after = Mm(0)
            para.add_run().add_picture(to_jpeg_buf(img), width=Inches(fit_w), height=Inches(fit_h))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── UI ─────────────────────────────────────────────────────────────────────────

st.title("Image to A4 Organizer")

with st.sidebar:
    st.header("Settings")

    # Compression
    st.subheader("Compression")
    enable_comp = st.checkbox("Enable Compression", value=True)
    comp_level = None
    if enable_comp:
        comp_level = st.radio(
            "Level",
            ["Low", "Medium", "High"],
            index=1,
            captions=["High quality, larger file", "Balanced", "Small file, lower quality"],
        )

    st.subheader("Metadata")
    strip_meta = st.checkbox("Strip metadata", value=True)

# Images per page
    st.subheader("Images per Page")
    ipp_choice = st.radio(
        "images_per_page",
        ["1", "2", "4", "Other"],
        index=2,
        label_visibility="collapsed",
    )
    if ipp_choice == "Other":
        ipp = int(
            st.number_input("Custom number", min_value=1, max_value=50, value=6, step=1)
        )
    else:
        ipp = int(ipp_choice)

# Export format
    st.subheader("Export Format")
    want_pdf = st.checkbox("PDF", value=True)
    want_html = st.checkbox("HTML (printable)", value=False)
    want_docx = st.checkbox("DOCX (Word)", value=False)

# Main upload area
uploaded = st.file_uploader(
    "Upload Images",
    type=["jpg", "jpeg", "png", "webp", "bmp", "tiff"],
    accept_multiple_files=True,
)

if uploaded:
    images = [Image.open(f) for f in uploaded]
    cols_shape, rows_shape = grid_shape(ipp)
    n_pages = math.ceil(len(images) / ipp)

    st.info(
        f"**{len(images)}** image(s)  →  **{n_pages}** page(s)  |  "
        f"Grid: **{cols_shape} × {rows_shape}**  |  "
        f"Compression: **{'None' if not enable_comp else comp_level}**"
    )

    with st.expander("Preview uploaded images", expanded=True):
        thumb_cols = st.columns(min(len(images), 5))
        for i, img in enumerate(images):
            thumb_cols[i % 5].image(
                img, caption=uploaded[i].name, use_container_width=True
            )

    if st.button("Generate", type="primary", use_container_width=True):
        if not want_pdf and not want_html and not want_docx:
            st.warning("Select at least one export format in the sidebar.")
        else:
            with st.spinner("Generating…"):
                st.session_state.pdf_data = (
                    generate_pdf(images, ipp, comp_level, strip_meta) if want_pdf else None
                )
                st.session_state.html_data = (
                    generate_html(images, ipp, comp_level, strip_meta) if want_html else None
                )
                st.session_state.docx_data = (
                    generate_docx(images, ipp, comp_level, strip_meta) if want_docx else None
                )
                st.session_state.generated = True

if st.session_state.get("generated"):
    st.success("Done! Download your files below.")
    active = [k for k in ("pdf_data", "html_data", "docx_data") if st.session_state.get(k)]
    dcols = st.columns(len(active)) if active else []
    col_iter = iter(dcols)
    if st.session_state.get("pdf_data"):
        next(col_iter).download_button(
            "⬇ Download PDF",
            st.session_state.pdf_data,
            file_name="images.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
    if st.session_state.get("html_data"):
        next(col_iter).download_button(
            "⬇ Download HTML",
            st.session_state.html_data,
            file_name="images.html",
            mime="text/html",
            use_container_width=True,
        )
    if st.session_state.get("docx_data"):
        next(col_iter).download_button(
            "⬇ Download DOCX",
            st.session_state.docx_data,
            file_name="images.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    if st.session_state.get("html_data"):
        with st.expander("Preview layout (HTML)"):
            st.components.v1.html(
                st.session_state.html_data, height=640, scrolling=True
            )
